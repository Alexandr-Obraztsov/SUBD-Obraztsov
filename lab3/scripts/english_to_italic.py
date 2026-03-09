#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Выделяет все английские слова (только латиница a-zA-Z) курсивом в документе Word.
Остальное форматирование (жирный, шрифт, размер, цвет и т.д.) сохраняется.
"""

import re
import sys
import argparse
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("Установите python-docx: pip install python-docx", file=sys.stderr)
    sys.exit(1)


def is_english_word(segment: str) -> bool:
    """Сегмент — английское слово, если состоит только из латинских букв."""
    return bool(segment and re.fullmatch(r"[a-zA-Z]+", segment))


def split_into_segments(text: str):
    """Разбивает текст на сегменты: латинские слова и всё остальное. Возвращает [(текст, is_english), ...]."""
    if not text:
        return []
    parts = re.findall(r"[a-zA-Z]+|[^a-zA-Z]+", text)
    return [(p, is_english_word(p)) for p in parts]


def get_run_format(run):
    """Снимает копию форматирования run (до удаления run из документа)."""
    fmt = {}
    fmt["bold"] = run.bold
    fmt["underline"] = run.underline
    fmt["italic"] = run.italic
    try:
        fmt["font_name"] = run.font.name
        fmt["font_size"] = run.font.size
        fmt["color_rgb"] = run.font.color.rgb if run.font.color.rgb is not None else None
    except Exception:
        fmt["font_name"] = fmt["font_size"] = fmt["color_rgb"] = None
    return fmt


def apply_run_format(fmt, target_run, italic_override=None):
    """Применяет сохранённое форматирование к target_run. italic_override задаёт курсив (если не None)."""
    target_run.bold = fmt.get("bold")
    target_run.underline = fmt.get("underline")
    target_run.italic = italic_override if italic_override is not None else fmt.get("italic")
    try:
        if fmt.get("font_name"):
            target_run.font.name = fmt["font_name"]
        if fmt.get("font_size"):
            target_run.font.size = fmt["font_size"]
        if fmt.get("color_rgb"):
            target_run.font.color.rgb = fmt["color_rgb"]
    except Exception:
        pass


def process_paragraph(paragraph):
    """
    В параграфе все английские слова делаются курсивом.
    Остальное форматирование каждого run сохраняется.
    """
    # Собираем по каждому run: (run, сегменты, копия формата)
    runs_data = []
    for run in paragraph.runs:
        text = run.text
        if not text:
            continue
        segments = split_into_segments(text)
        if not segments:
            continue
        fmt = get_run_format(run)
        runs_data.append((run, segments, fmt))

    if not runs_data:
        return

    # Удаляем все участвовавшие run'ы
    for run, _, _ in runs_data:
        run._r.getparent().remove(run._r)

    # Добавляем новые run'ы в том же порядке, с сохранённым форматом
    for _run, segments, fmt in runs_data:
        for text, is_english in segments:
            new_run = paragraph.add_run(text)
            apply_run_format(fmt, new_run, italic_override=is_english)


def process_document(doc):
    """Обрабатывает весь документ: параграфы, таблицы, колонтитулы."""
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            process_paragraph(paragraph)
        for paragraph in section.footer.paragraphs:
            process_paragraph(paragraph)


def main():
    parser = argparse.ArgumentParser(
        description="Выделяет английские слова курсивом в документе .docx, остальное форматирование не меняет."
    )
    parser.add_argument(
        "path",
        nargs="?",
        default=None,
        help="Путь к файлу .docx (по умолчанию: lab2/Отчет.docx относительно скрипта)",
    )
    args = parser.parse_args()

    if args.path:
        doc_path = Path(args.path)
    else:
        doc_path = Path(__file__).resolve().parent.parent / "Отчет.docx"

    if not doc_path.is_file():
        print(f"Файл не найден: {doc_path}", file=sys.stderr)
        sys.exit(1)

    doc = Document(str(doc_path))
    process_document(doc)
    doc.save(str(doc_path))
    print(f"Готово: английские слова выделены курсивом в {doc_path}")


if __name__ == "__main__":
    main()
